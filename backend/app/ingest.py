"""
Document ingestion: load → clean → chunk → embed → store.

Supported formats: PDF, TXT, DOCX, CSV, Markdown, HTML
Edge cases handled: empty files, corrupt files, oversized files, scanned PDFs (OCR fallback).
"""
import csv
import io
import logging
from dataclasses import dataclass

import pdfplumber
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fastapi import UploadFile
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.database import SessionLocal
from app.llm import get_embeddings
from app.models import ChunkMetadata, Document
from app.vectorstore import Chunk, FAISSVectorStore

logger = logging.getLogger(__name__)

# ── Constants ─────────────────────────────────────────────────────────────────

CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
EMBED_BATCH = 100
MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB hard limit

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".csv", ".md", ".html", ".htm"}

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    length_function=len,
)


# ── Parsed page ───────────────────────────────────────────────────────────────

@dataclass
class ParsedPage:
    text: str
    page: int


# ── Format parsers ────────────────────────────────────────────────────────────

def _parse_pdf(data: bytes) -> list[ParsedPage]:
    pages: list[ParsedPage] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(ParsedPage(text=text, page=i + 1))
    except Exception as exc:
        raise ValueError(f"Could not parse PDF: {exc}") from exc

    if not pages:
        # Scanned PDF — all pages returned no text
        pages = _ocr_fallback(data)

    return pages


def _ocr_fallback(data: bytes) -> list[ParsedPage]:
    """Attempt pytesseract OCR; flag gracefully if unavailable."""
    try:
        import pytesseract
        from PIL import Image as PILImage
        import pdf2image  # type: ignore

        images = pdf2image.convert_from_bytes(data, dpi=200)
        pages: list[ParsedPage] = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img) or ""
            if text.strip():
                pages.append(ParsedPage(text=text, page=i + 1))
        if pages:
            return pages
        raise ValueError("OCR found no readable text in this scanned PDF.")
    except ImportError:
        raise ValueError(
            "This PDF appears to be scanned (no embedded text). "
            "Install pytesseract + pdf2image + poppler for OCR support, "
            "or upload a text-based PDF."
        )


def _parse_txt(data: bytes) -> list[ParsedPage]:
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception as exc:
        raise ValueError(f"Could not decode text file: {exc}") from exc
    return [ParsedPage(text=text, page=1)]


def _parse_docx(data: bytes) -> list[ParsedPage]:
    try:
        doc = DocxDocument(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        if not text.strip():
            raise ValueError("DOCX file contains no readable text.")
        return [ParsedPage(text=text, page=1)]
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not parse DOCX: {exc}") from exc


def _parse_csv(data: bytes) -> list[ParsedPage]:
    try:
        text_io = io.StringIO(data.decode("utf-8", errors="replace"))
        reader = csv.reader(text_io)
        rows = list(reader)
        if not rows:
            raise ValueError("CSV file is empty.")
        lines = [", ".join(row) for row in rows]
        text = "\n".join(lines)
        return [ParsedPage(text=text, page=1)]
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not parse CSV: {exc}") from exc


def _parse_markdown(data: bytes) -> list[ParsedPage]:
    try:
        text = data.decode("utf-8", errors="replace")
        if not text.strip():
            raise ValueError("Markdown file is empty.")
        return [ParsedPage(text=text, page=1)]
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not parse Markdown: {exc}") from exc


def _parse_html(data: bytes) -> list[ParsedPage]:
    try:
        soup = BeautifulSoup(data, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "head"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        if not text.strip():
            raise ValueError("HTML file contains no readable text.")
        return [ParsedPage(text=text, page=1)]
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not parse HTML: {exc}") from exc


# ── Dispatcher ────────────────────────────────────────────────────────────────

def _detect_and_parse(filename: str, content_type: str, data: bytes) -> list[ParsedPage]:
    name = (filename or "").lower()
    mime = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in mime:
        return _parse_pdf(data)
    if name.endswith(".docx") or "wordprocessingml" in mime or "msword" in mime:
        return _parse_docx(data)
    if name.endswith(".csv") or "csv" in mime:
        return _parse_csv(data)
    if name.endswith((".md", ".markdown")) or "markdown" in mime:
        return _parse_markdown(data)
    if name.endswith((".html", ".htm")) or "html" in mime:
        return _parse_html(data)
    if name.endswith(".txt") or mime.startswith("text/"):
        return _parse_txt(data)

    ext = "." + name.rsplit(".", 1)[-1] if "." in name else ""
    raise ValueError(
        f"Unsupported file type '{filename}'. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


# ── Public API ────────────────────────────────────────────────────────────────

async def ingest_file(file: UploadFile, store: FAISSVectorStore) -> int:
    """Parse → chunk → embed → store. Returns the new document_id."""
    data = await file.read()
    filename = file.filename or "upload"
    content_type = file.content_type or ""

    # Validate
    if not data:
        raise ValueError("Uploaded file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(
            f"File exceeds the 50 MB limit "
            f"({len(data) / 1024 / 1024:.1f} MB). "
            "Split it into smaller files."
        )

    pages = _detect_and_parse(filename, content_type, data)
    if not pages:
        raise ValueError("No text could be extracted from the file.")

    # Persist document record
    with SessionLocal() as db:
        doc = Document(filename=filename, content_type=content_type)
        db.add(doc)
        db.commit()
        db.refresh(doc)
        document_id = doc.id

    # Build raw chunks with metadata
    raw: list[dict] = []
    for p in pages:
        for i, text in enumerate(_splitter.split_text(p.text)):
            raw.append({
                "text": text,
                "metadata": {
                    "document_id": document_id,
                    "filename": filename,
                    "page": p.page,
                    "chunk_index": i,
                },
            })

    if not raw:
        raise ValueError("Document produced no usable text chunks after splitting.")

    # Embed in batches (prevents OOM on large docs)
    embedder = get_embeddings()
    chunks: list[Chunk] = []
    for start in range(0, len(raw), EMBED_BATCH):
        batch = raw[start: start + EMBED_BATCH]
        embeddings = embedder.embed_documents([r["text"] for r in batch])
        for r, emb in zip(batch, embeddings):
            chunks.append(Chunk(text=r["text"], metadata=r["metadata"], embedding=emb))

    store.add(chunks)

    # Persist chunks to SQLite + update chunk count (atomic)
    with SessionLocal() as db:
        # Remove any stale chunks for this document (idempotent re-index)
        db.query(ChunkMetadata).filter(ChunkMetadata.document_id == document_id).delete()

        for chunk in chunks:
            m = chunk.metadata
            db.add(ChunkMetadata(
                document_id=m["document_id"],
                chunk_index=m["chunk_index"],
                page=m["page"],
                text=chunk.text,
            ))

        doc = db.get(Document, document_id)
        if doc:
            doc.chunk_count = len(chunks)
        db.commit()

    logger.info("Ingested '%s': %d pages, %d chunks", filename, len(pages), len(chunks))
    return document_id
